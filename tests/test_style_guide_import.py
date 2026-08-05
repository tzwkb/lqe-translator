import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from lqe_io import _load_style_guide


def add_numbered_paragraph(doc: Document, text: str, num_id: int = 5) -> None:
    para = doc.add_paragraph(text)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    para._p.get_or_add_pPr().append(num_pr)


class StyleGuideImportTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.root = Path(self.tempdir.name)

    def tearDown(self):
        self.tempdir.cleanup()

    def test_docx_preserves_block_order_tables_and_numbering(self):
        path = self.root / "guide.docx"
        doc = Document()
        doc.add_heading("Rules", level=1)
        add_numbered_paragraph(doc, "First rule")
        add_numbered_paragraph(doc, "Second rule")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Source"
        table.cell(0, 1).text = "Target"
        table.cell(1, 0).text = "A|B"
        table.cell(1, 1).text = "Line one\nLine two"
        doc.add_paragraph("After table")
        doc.save(path)

        text = _load_style_guide(str(path))

        expected = [
            "# Rules",
            "1. First rule",
            "2. Second rule",
            "| Source | Target |",
            "| --- | --- |",
            r"| A\|B | Line one / Line two |",
            "After table",
        ]
        positions = [text.index(item) for item in expected]
        self.assertEqual(positions, sorted(positions))

    def test_paragraph_only_docx_keeps_legacy_layout(self):
        path = self.root / "legacy.docx"
        doc = Document()
        doc.add_paragraph("Title")
        doc.add_heading("Rules", level=1)
        doc.add_paragraph("Body")
        doc.save(path)

        self.assertEqual(_load_style_guide(str(path)), "Title\n\n# Rules\nBody")


if __name__ == "__main__":
    unittest.main()
